import os
import pdfplumber


def extract_jd_from_document(file_path):
    """Extract job description text from a PDF or DOCX file.

    Args:
        file_path: Path to the uploaded document

    Returns:
        Extracted text string
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF, DOCX, or TXT.")


def _extract_from_pdf(file_path):
    """Extract text from PDF."""
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    if not text.strip():
        raise ValueError("Could not extract text from PDF. The file may be image-based or empty.")
    return text


def _extract_from_docx(file_path):
    """Extract text from DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to read DOCX files. "
            "Install it: pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables (some JDs are in table format)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    text = "\n".join(paragraphs)
    if not text.strip():
        raise ValueError("Could not extract text from DOCX. The file may be empty.")
    return text
